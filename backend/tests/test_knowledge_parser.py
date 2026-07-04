"""knowledge/parser.py 单元测试（P3-13 PDF/DOCX 真实解析）。

覆盖：
- text/plain / text/markdown 走 UTF-8 严格 decode
- application/pdf 经 pypdf 提取文本（构造最小 PDF）
- DOCX 经 python-docx 提取段落（构造最小 DOCX）
- 提取失败（非法 UTF-8 / 空 PDF / 空 DOCX / 不支持类型 / 超大）抛 ValidationError
"""

from __future__ import annotations

import io

import pytest

from app.core.exceptions import ValidationError
from app.domains.knowledge import parser
from app.domains.knowledge.parser import DOCX_MIME, extract_text

# ===================== 文本类 =====================


def test_extract_text_plain_decodes_utf8() -> None:
    """text/plain 正常 UTF-8 decode。"""
    content = "AIOps 是 AI 原生运营控制台。".encode()
    assert extract_text(content, "text/plain") == "AIOps 是 AI 原生运营控制台。"


def test_extract_text_markdown_decodes_utf8() -> None:
    """text/markdown 走同一条 UTF-8 decode 路径。"""
    content = b"# Title\n\nsome **markdown** body"
    assert extract_text(content, "text/markdown") == "# Title\n\nsome **markdown** body"


def test_extract_text_invalid_utf8_raises() -> None:
    """非法 UTF-8 字节必须抛 ValidationError，而非静默替换为 U+FFFD。"""
    # 0xff 在 UTF-8 中是非法起始字节
    with pytest.raises(ValidationError) as exc_info:
        extract_text(b"\xff\xfe illegal bytes", "text/plain")
    assert "UTF-8" in str(exc_info.value)


# ===================== PDF =====================


def _make_minimal_pdf(text: str = "Hello PDF test") -> bytes:
    """构造包含指定 ASCII 文本的最小单页 PDF（Helvetica）。

    手工拼装对象 + 计算 xref 偏移，确保 pypdf 可解析并 extract_text 命中文本。
    """
    content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]"
            b" /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>"
        ),
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    pdf = b"%PDF-1.4\n"
    offsets: list[int] = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_offset = len(pdf)
    pdf += b"xref\n"
    pdf += f"0 {len(objects) + 1}\n".encode()
    pdf += b"0000000000 65535 f \n"
    for off in offsets:
        pdf += f"{off:010d} 00000 n \n".encode()
    pdf += b"trailer\n"
    pdf += f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode()
    pdf += b"startxref\n"
    pdf += f"{xref_offset}\n".encode()
    pdf += b"%%EOF"
    return pdf


def _make_empty_pdf() -> bytes:
    """构造无文本内容的空白单页 PDF（pypdf PdfWriter 生成）。"""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def test_extract_pdf_extracts_text() -> None:
    """PDF 提取能命中文本内容（P3-13 核心路径）。"""
    pdf = _make_minimal_pdf("Hello PDF test")
    text = extract_text(pdf, "application/pdf")
    assert "Hello PDF test" in text


def test_extract_pdf_empty_raises() -> None:
    """空白 PDF 提取不到文本应抛 ValidationError（避免空 embedding）。"""
    with pytest.raises(ValidationError) as exc_info:
        extract_text(_make_empty_pdf(), "application/pdf")
    assert "PDF" in str(exc_info.value)


def test_extract_pdf_corrupt_raises() -> None:
    """损坏的 PDF 字节应抛 ValidationError 而非崩成 U+FFFD。"""
    with pytest.raises(ValidationError):
        extract_text(b"not a real pdf %PDF-1.4 garbage", "application/pdf")


# ===================== DOCX =====================


def _make_minimal_docx(text: str = "Hello DOCX test") -> bytes:
    """用 python-docx 构造包含一段文本的最小 DOCX。"""
    from docx import Document

    document = Document()
    document.add_paragraph(text)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _make_empty_docx() -> bytes:
    """构造无段落文本的空 DOCX。"""
    from docx import Document

    document = Document()
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_extract_docx_extracts_text() -> None:
    """DOCX 提取能命中段落文本（P3-13 核心路径）。"""
    docx = _make_minimal_docx("Hello DOCX test")
    text = extract_text(docx, DOCX_MIME)
    assert "Hello DOCX test" in text


def test_extract_docx_multiple_paragraphs() -> None:
    """多段落 DOCX 提取后以换行连接。"""
    from docx import Document

    document = Document()
    document.add_paragraph("first paragraph")
    document.add_paragraph("second paragraph")
    buf = io.BytesIO()
    document.save(buf)
    text = extract_text(buf.getvalue(), DOCX_MIME)
    assert "first paragraph" in text
    assert "second paragraph" in text
    assert "\n" in text


def test_extract_docx_empty_raises() -> None:
    """空 DOCX 提取不到文本应抛 ValidationError。"""
    with pytest.raises(ValidationError) as exc_info:
        extract_text(_make_empty_docx(), DOCX_MIME)
    assert "DOCX" in str(exc_info.value)


def test_extract_docx_corrupt_raises() -> None:
    """损坏的 DOCX 字节应抛 ValidationError。"""
    with pytest.raises(ValidationError):
        extract_text(b"not a real docx (zip) bytes", DOCX_MIME)


# ===================== 其他 / 边界 =====================


def test_extract_unsupported_type_raises() -> None:
    """非白名单 content-type 应抛 ValidationError。"""
    with pytest.raises(ValidationError) as exc_info:
        extract_text(b"whatever", "application/octet-stream")
    assert "octet-stream" in str(exc_info.value)


def test_extract_pdf_oversized_raises(monkeypatch: pytest.MonkeyPatch) -> None:
    """PDF 提取前校验原始字节大小，超限抛 ValidationError。"""
    # monkeypatch 到小阈值，避免在测试中分配 50MB 字节串
    monkeypatch.setattr(parser, "MAX_DOC_BYTES", 100)
    with pytest.raises(ValidationError) as exc_info:
        extract_text(b"x" * 101, "application/pdf")
    assert "超" in str(exc_info.value) or "MB" in str(exc_info.value)


def test_extract_docx_oversized_raises(monkeypatch: pytest.MonkeyPatch) -> None:
    """DOCX 提取前校验原始字节大小，超限抛 ValidationError。"""
    monkeypatch.setattr(parser, "MAX_DOC_BYTES", 100)
    with pytest.raises(ValidationError):
        extract_text(b"x" * 101, DOCX_MIME)


def test_extract_text_plain_not_size_checked() -> None:
    """文本类不经过 _check_size（沿用 service 层 decoded size 校验）。

    即便字节数大于 monkeypatch 阈值也不应被 parser 拦截——此处验证默认
    MAX_DOC_BYTES 下文本路径不触发 _check_size（50MB 远大于测试内容）。
    """
    # 不 monkeypatch：默认 50MB，文本路径不应触发 _check_size
    text = extract_text("普通文本".encode(), "text/plain")
    assert text == "普通文本"
